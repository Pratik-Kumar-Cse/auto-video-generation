import os
from docx import Document
from typing import Annotated
from autogen.agentchat import Agent
import json

directory = "./content/"

if not os.path.exists(directory):
    os.makedirs(directory)


def save_docx(content: Annotated[Agent, "Agent"]):
    content = Agent.last_message()["content"]
    existing_files = [file for file in os.listdir(directory) if file.endswith(".docx")]
    count = len(existing_files) + 1
    file_name = f"Script{count}.docx"
    file_path = os.path.join(directory, file_name)
    doc = Document()
    doc.add_paragraph(content)
    doc.save(file_path)


def save_json(content):
    json_data = {"script": content}
    existing_files = [file for file in os.listdir(directory) if file.endswith(".json")]
    count = len(existing_files) + 1
    file_name = f"script.json"
    file_path = os.path.join(directory, file_name)
    with open(file_path, "w") as json_file:
        json.dump(json_data, json_file, indent=4)


def save_docx_reel(Agent: Annotated[Agent, "Agent"]):
    content = Agent.last_message()["content"]
    existing_files = [file for file in os.listdir(directory) if file.endswith(".docx")]
    count = len(existing_files) + 1
    file_name = f"Script{count}.docx"
    file_path = os.path.join(directory, file_name)
    doc = Document()
    doc.add_paragraph(content)
    doc.save(file_path)


def save_json_reel(content):
    json_data = {"script": content}
    existing_files = [file for file in os.listdir(directory) if file.endswith(".json")]
    count = len(existing_files) + 1
    file_name = f"script.json"
    file_path = os.path.join(directory, file_name)
    with open(file_path, "w") as json_file:
        json.dump(json_data, json_file, indent=4)


def store_script_title(search_terms):
    """
    Store the script and story title points in a JSON file.
    Args:
        script (str): The script content.
        title_points (list): A list of title points.
        filename (str): The name of the JSON file to store the data.
    """
    data = {"data": json.loads(search_terms)}
    file_name = f"search_terms.json"
    file_path = os.path.join(directory, file_name)
    with open(file_path, "w") as json_file:
        json.dump(data, json_file)


def store_clip_data(data):
    """
    Store the script and story title points in a JSON file.
    Args:
        script (str): The script content.
        title_points (list): A list of title points.
        filename (str): The name of the JSON file to store the data.
    """
    data = {data: data}
    file_name = "clips.json"
    file_path = os.path.join(directory, file_name)
    with open(file_path, "w") as json_file:
        json.dump(data, json_file)

    return "success"


def store_image_data(data):
    """
    Store the script and story title points in a JSON file.
    Args:
        script (str): The script content.
        title_points (list): A list of title points.
        filename (str): The name of the JSON file to store the data.
    """
    data = {data: data}
    file_name = "images.json"
    file_path = os.path.join(directory, file_name)
    with open(file_path, "w") as json_file:
        json.dump(data, json_file)

    return "success"


def store_image_link(data: object):
    """
    Store the script and story title points in a JSON file.
    Args:
        script (str): The script content.
        title_points (list): A list of title points.
        filename (str): The name of the JSON file to store the data.
    """
    data = {"data": json.loads(data)}
    file_name = "image_links.json"
    file_path = os.path.join(directory, file_name)
    with open(file_path, "w") as json_file:
        json.dump(data, json_file)

    return "success"
